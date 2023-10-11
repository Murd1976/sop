from langchain.llms import OpenAI
#from langchain.chat_models import ChatOpenAI
from langchain.chains import LLMRequestsChain, LLMChain
from langchain.chains.summarize import load_summarize_chain
from langchain.chains.question_answering import load_qa_chain
#from langchain.chains import AnalyzeDocumentChain #, CondenseChain
from langchain.document_loaders import TextLoader
from langchain.docstore.document import Document
#import requests
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import Chroma, FAISS
#from langchain.text_splitter import CharacterTextSplitter
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
import docx

import os, shutil
import re
import time
from datetime import datetime

import openai
import tiktoken
import copy
#from dotenv import load_dotenv
from .web_utils import scrape_site

openai.api_key = os.getenv('OPENAI_API_KEY')

class WorkerОpenAI_SOP():
  sys_dir = 'ai_sop/db/'
  # папка базы данных Proposal
  persist_directory = ''
  # путь к учебным материалам Proposal
  data_directory = ''
  #db_proposal_file = 'train_data_jobs.jsonl'
  system_directory = ''
  system_doc_file = ''
  web_directory = ''
  
  system_doc_web_prop = 'html_proposal_instruction.txt'
  __embedding_new =  False

  def __init__(self, em_framework = 'chroma', system_file = '', company_dir = None, mod = 'gpt-3.5-turbo-0301'):
    self.model = mod
    self.em_framework = em_framework
    self.debug_log = []
    self.persist_directory = self.sys_dir + 'sop/embedding/' + company_dir
    self.system_directory = self.sys_dir + 'sop/sys/' + company_dir
    self.system_doc_file = system_file
    self.data_directory = self.sys_dir + 'data/' + company_dir
    self.web_directory = self.sys_dir + 'sop/web/' + company_dir
    
    if not os.path.exists(self.persist_directory):
        os.mkdir(self.persist_directory)
    if not os.path.exists(self.system_directory):
        os.mkdir(self.system_directory)
    if not os.path.exists(self.data_directory):
        os.mkdir(self.data_directory)
    if not os.path.exists(self.web_directory):
        os.mkdir(self.web_directory)
    
        
    self.chat_manager_system = self.load_file(self.system_directory, self.system_doc_file)  
    if em_framework == 'chroma':
        self.persist_directory = self.persist_directory + 'chroma/'
        if not os.path.exists(self.persist_directory):
            os.mkdir(self.persist_directory)
        # Если База данных embedding уже создана ранее
        if os.path.exists(self.persist_directory + 'embedding_info.inf'):
            print("We use a ready-made database. Path: ", self.persist_directory)
            self.debug_log.append("We use a ready-made database. Path: " + self.persist_directory)
            self.db = Chroma(persist_directory=self.persist_directory,
                                embedding_function=OpenAIEmbeddings())
        else:
            print("Embeddings database not found!")
            self.debug_log.append("Embeddings database not found!")
            self.db = Chroma(persist_directory=self.persist_directory, embedding_function=OpenAIEmbeddings())

    if em_framework == 'faiss':
        self.persist_directory = self.persist_directory + 'faiss/'
        # Load index from disk
        self.db = FAISS.load_local(self.persist_directory + 'faiss_index')

  # def get_key(self):
  #   openai.api_key = getpass.getpass(prompt='Введите секретный ключ для сервиса chatGPT: ')
  #   import os
  #   os.environ["OPENAI_API_KEY"] = openai.api_key

  @property  
  def em_mode(self):
    return self.__embedding_new
  
  @em_mode.setter  
  def em_mode(self, mode):
    if mode == '1':
        self.__embedding_new = False
    else:
        self.__embedding_new = True

  def load_document_text(self, url: str) -> str:
      with open(url, "r") as f:
        text = f.read()

      return text
      
  def file_list(self, f_path):
                
    f_list = os.listdir(f_path)
    f_list = [ff for ff in f_list if (os.path.isfile(f_path + ff) and ('~' not in ff))]
        
    return f_list

  def del_all_in_dir(self, dir_path):
    if os.path.exists(dir_path):
        for files in os.listdir(dir_path):
            path = os.path.join(dir_path, files)
            try:
                shutil.rmtree(path)
            except OSError:
                os.remove(path)
                
  def get_text_from(self, doc_dir, file_):
  
    with open(doc_dir + file_, "r") as f:                
        self.debug_log.append("File is loading:" + file_)
        buf = file_.split('.')
        if buf[-1] in ['doc', 'docx'] :
            # Load DOC/DOCX document
            document_txt = ""
            try:
                doc = docx.Document(self.data_directory + file_)
                for docpara in doc.paragraphs:
                    document_txt += docpara.text + '\n'
            except Exception as e:
                self.debug_log.append(f"Error processing file {file_}: {str(e)}")
                #print(document_txt)
        elif buf[-1] == 'txt':
            # This is a long document we can split up.
            document_txt = self.load_document_text(self.data_directory + file_)
            
    return document_txt

  def num_tokens_from_string(self, string: str, encoding_name: str) -> int:
      """Returns the number of tokens in a text string."""
      encoding = tiktoken.get_encoding(encoding_name)
      num_tokens = len(encoding.encode(string))
      return num_tokens
      
      
    
  def create_embedding(self, doc_dir=""):
    if self.em_framework == 'chroma':
        self.create_embedding_chroma(doc_dir= doc_dir)
    elif self.em_framework == 'faiss':    
        self.create_embedding_faiss(doc_dir= doc_dir)
        
  def create_embedding_chroma(self, doc_dir=""):
    

    self.source_chunks = []
    self.buf_chunks = []
    splitter = RecursiveCharacterTextSplitter(["<Chunk>", '\n\n', '\n', ' '], chunk_size=1024, chunk_overlap=0)

    
    #print('Files: ', os.listdir(doc_dir))

    #print("File is loading: ", doc_dir)
    #print(f'Chroma dir: {persist_directory}')
    
    # проходимся по всем данным
    count_em_token = 0 # all tokens gones thrue embeddings
    f_info = open(self.persist_directory + 'embedding_info.inf', 'w')
    f_info.write('Used files: \n')
    for file_ in sorted(self.file_list(doc_dir)): #go thrue all files in doc_dir
        
        
        print("Загружается файл: ", file_)
        
        self.debug_log.append("File is loading:" + doc_dir)
        # разбиваем на несколько частей с помощью метода split_text
        if (os.path.isfile(doc_dir + file_)):
            
            try:
                document_txt = self.get_text_from(doc_dir, file_) #get text from current file
                f_info.write(str(file_) + '\n')
            except Exception as e:
                self.debug_log.append(f"Error processing file {file_}: {str(e)}") 
                
            count_token = 0 #counter for OpenAI (tokens limit 140K per minute )    
            for chunk in splitter.split_text(document_txt):
                #print('Длина символов =  ', len(chunk))

                count_token += self.num_tokens_from_string(chunk, "cl100k_base")
                if count_token > 140000:
                   
                    #print('Count: ', count_token, ' Tokens:  ', num_tokens_from_string(' '.join([x.page_content for x in self.buf_chunks]), "cl100k_base"))
                  
                    count_token = 0
                    self.source_chunks.append(copy.deepcopy(self.buf_chunks))
                   
                    #print('Size: ', len(self.buf_chunks), '\n')
                    self.buf_chunks.clear()
                   
                    self.buf_chunks.append(Document(page_content=chunk, metadata={'source': file_}))
                else:
                    self.buf_chunks.append(Document(page_content=chunk, metadata={'source': file_}))
                        
            self.source_chunks.append(copy.deepcopy(self.buf_chunks))
                
            if self.__embedding_new :
                print('Creating new embeddings!')
                self.del_all_in_dir(self.persist_directory)
                self.debug_log.append('Previouse embeddings deleted!')
                self.db = Chroma(persist_directory=self.persist_directory, embedding_function=OpenAIEmbeddings())
                    
            else:
                print('Adding embeddings!')
                   
            for i in range(len(self.source_chunks)):
                self.db.add_documents(documents=self.source_chunks[i])
                print('sSize: ', len(self.source_chunks[i]))
                self.debug_log.append('sSize: ' + str(len(self.source_chunks[i])))
                count_token = self.num_tokens_from_string(' '.join([x.page_content for x in self.source_chunks[i]]), "cl100k_base")
                count_em_token += count_token
                print(f'Curr 140K chunk: {i} Counter: Global- {count_em_token} Curr file- {count_token} ')
                self.debug_log.append(str(i) + 'Counter: ' + str(count_token))
                time.sleep(77)
                  
            self.db.persist()
            self.source_chunks.clear()
                
    #print('Count: ', count_token, ' Tokens:  ', num_tokens_from_string(' '.join([x.page_content for x in buf_chunks]), "cl100k_base"))
    #print('Size: ', len(buf_chunks), '\n')
    
    #print('G_Size: ', self.source_chunks[0])
    #print('G_Size: ', self.source_chunks[1])
    #print('G_Size: ', self.source_chunks[2])
       
    f_info.close()
    
    print('\n ===========================================:')
    print('\n Number of tokens in source document: ', count_em_token)
    print('\n Request price: ', 0.0004*(count_em_token/1000), ' $')
    print('\n ===========================================')
    
    self.debug_log.append('\n ===========================================: ')
    self.debug_log.append('\n Number of tokens in source document: ' + str(count_em_token))
    self.debug_log.append('\n Request price: ' + str(0.0004*(count_em_token/1000)) + ' $')
    self.debug_log.append('\n =========================================== \n')


  def insert_newlines(self, text: str, max_len: int = 170) -> str:
      words = text.split()
      lines = []
      current_line = ""
      for word in words:
          if len(current_line + " " + word) > max_len:
              lines.append(current_line)
              current_line = ""
          current_line += " " + word
      lines.append(current_line)
      return "\n".join(lines)
      
  # пример подсчета токенов
  def num_tokens_from_messages(self, messages):
      """Returns the number of tokens used by a list of messages."""
      try:
          encoding = tiktoken.encoding_for_model(self.model)
      except KeyError:
          encoding = tiktoken.get_encoding("cl100k_base")
      if self.model == "gpt-3.5-turbo-0301":  # note: future models may deviate from this
          num_tokens = 0
          for message in messages:
              num_tokens += 4  # every message follows <im_start>{role/name}\n{content}<im_end>\n
              for key, value in message.items():
                  num_tokens += len(encoding.encode(value))
                  if key == "name":  # if there's a name, the role is omitted
                      num_tokens += -1  # role is always required and always 1 token
          num_tokens += 2  # every reply is primed with <im_start>assistant
          return num_tokens
      else:
          raise NotImplementedError(f"""num_tokens_from_messages() is not presently implemented for model {model}.
  See https://github.com/openai/openai-python/blob/main/chatml.md for information on how messages are converted to tokens.""")
  
  
  def web_text_filter_2(self, f_name):
    output_file = self.web_directory + 'out_' + f_name
    # Load text from .txt file
    loader = TextLoader(self.web_directory + f_name)
    doc = loader.load()
    #print(doc)
    # Split large text file into chunks
    splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(chunk_size=2048)
    texts = splitter.split_documents([doc[0]])
    
    print(texts[0])
    count_token = 0
    for chunk in texts:
        
        count_token += self.num_tokens_from_string(chunk.page_content, "cl100k_base")
        print(count_token)
    
    # Initialize the chain 
    combine_chain = CondenseChain()
    llm = OpenAI(temperature=0)
    analyzer = AnalyzeDocumentChain(llm=llm, combine_docs_chain=combine_chain)

    # Pass splits through chain to analyze and clean
    results = analyzer.run(documents=texts)
    print(results)
    with open(output_file, 'w', encoding='utf-8') as file:
        # Extract cleaned text 
        res_text = [r.cleaned_text for r in results]
        file.write(res_text)
    
    return res_text
  
  def web_text_filter(self, f_name):
    def generate_blog_post(doc_content):
        #docs = search_index.similarity_search(topic, k=4)
        #inputs = [{"context": doc.page_content, "topic": topic} for doc in docs]
        inputs = [{"context": doc_content}]
        return chain.apply(inputs)
        
    output_file = self.web_directory + 'out_' + f_name
    
    # This is a long document we can split up.
    document_txt = self.load_document_text(self.web_directory + f_name)
    text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(chunk_size=2048, chunk_overlap=0)
    texts = text_splitter.split_text(document_txt)
    
    prompt_template_web = """The context below is the downloaded text from the WEB page, 
    you need to convert it into structured information, remove duplicate information and remove unnecessary text.:
    Context: {context}
    
    Analyzed data:"""

    PROMPT = PromptTemplate(template=prompt_template_web, input_variables=["context"])

    llm = OpenAI(temperature=0)

    chain = LLMChain(llm=llm, prompt=PROMPT)
    
    self.source_chunks = []
    self.buf_chunks = []
    count_token = 0
    res_text = ""
    with open(output_file, 'w', encoding='utf-8') as file:
        for chunk in texts:
            count_token += self.num_tokens_from_string(chunk, "cl100k_base")
            print(count_token)
            res_text = generate_blog_post(chunk)[0]["text"]
            file.write(res_text)
    res_text = generate_blog_post(texts[0])[0]["text"]
    #print(res_text)
    count_token = self.num_tokens_from_string(res_text, "cl100k_base")
    print(count_token)
    return res_text
    
  def load_file(self, f_dir, f_name):
    
    text = ""
    try:
        with open(f_dir + f_name, "r") as f:
            text = f.read()
        f.close()
        print("File is loading:" + f_dir + f_name)
        self.debug_log.append("File is loading:" + f_dir + f_name)
    except:
        
        print("File not found:" + f_dir + f_name)
        self.debug_log.append("File not found:" + f_dir + f_name)
        
    return text
    
  def load_analyzed_table(self, f_dir, f_name):
    res_list = []
    try:
        with open(f_dir + f_name, "r") as f:
            text = f.read()
        f.close()
        for line in text.splitlines():
            res_list.append(line)
        print("File is loading:" + f_dir + f_name)
        self.debug_log.append("File is loading:" + f_dir + f_name)
    except:
        res_list.append(f_name)
        print("File not found:" + f_dir + f_name)
        self.debug_log.append("File not found:" + f_dir + f_name)
        
    return res_list
    
  def doc_analyzer_fields(self):
    def analyze_chunk(doc_content, exists_list):
        #docs = search_index.similarity_search(topic, k=4)
        #inputs = [{"context": doc.page_content, "topic": topic} for doc in docs]
        inputs = [{"context": doc_content, "categories":exists_list}]
        return chain.apply(inputs)
        
    output_file = self.system_directory + 'categories.sop'
    
    analyzed_list = self.load_analyzed_table(self.data_directory, 'analyzed_cat.sop')
    
    exists_list = self.load_file(self.system_directory, 'categories.sop')
    
    #texts = []
    
    prompt_template_categories = """The context below is unprocessed information related to the activities of the company. 
    This information may include documents, company policies, downloaded content from the company website and online store, 
    search results. This information should be analyzed, identified and listed the main categories or topics and section that arise.
    Expected response:
    Product information
    Return policy
    Customer Requests
    Agent Answers
    Technical support
    Company History ... [and so on]
    
    Based on the generated categories, topics or section names, a knowledge base will be built, according to the principle of Wikipedia. 
    Some categories, topics and section names, may be subsections of existing ones. The result should be grouped as a "tree" and returned in .json format.
    
    In the Categories section, the result of previous requests with a similar task is presented  in .json format, 
    depending on the current result of the analysis, this list needs to be updated, supplemented or modified.
    
    Context: {context}
    
    Categories: {categories}
    
    Analysis result:"""

    PROMPT = PromptTemplate(template=prompt_template_categories, input_variables=["context", "categories"])

    llm = OpenAI(temperature=0)

    chain = LLMChain(llm=llm, prompt=PROMPT)
        
    count_token = 0
    res_text = ""
    
    text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(chunk_size=2048, chunk_overlap=0)
    for file_ in sorted(os.listdir(self.data_directory)):
        if (file_ not in analyzed_list):
            #print(analyzed_list)
            print("Загружается файл: ", file_)
            self.debug_log.append("File is loading:" + file_)
            buf = file_.split('.')
            if buf[-1] in ['doc', 'docx'] :
                # Load DOC/DOCX document
                document_txt = ""
                doc = docx.Document(self.data_directory + file_)
                for docpara in doc.paragraphs:
                    document_txt += docpara.text + '\n'
                #print(document_txt)
            else:
                # This is a long document we can split up.
                document_txt = self.load_document_text(self.data_directory + file_)
                
            with open(output_file, 'w', encoding='utf-8') as file:
                for chunk in text_splitter.split_text(document_txt):
                    count_token += self.num_tokens_from_string(str(chunk) + exists_list, "cl100k_base")
                    print(count_token)
                    res_text =analyze_chunk(chunk, exists_list)[0]["text"]
                    exists_list = res_text
                file.write(res_text)
            analyzed_list.append(file_)
            #print(document_txt)
            #texts.append(text_splitter.split_text(document_txt))
    
            #Сораняем игнор лист (файлы которые уже проанализированы )
            with open(self.data_directory + 'analyzed_cat.sop', 'w', encoding='utf-8') as file:
                for  line in analyzed_list:
                    file.write(line + '\n')
            
    #res_text = analyze_chunkt(texts[0])[0]["text"]
    print(res_text)
    #count_token = self.num_tokens_from_string(res_text, "cl100k_base")
    #print(count_token)
    return res_text
    
  def doc_analyzer_sop(self, file_, templ, instructions = ''):
    def analyze_chunk(doc_content, instructions = ''):
        #docs = search_index.similarity_search(topic, k=4)
        #inputs = [{"context": doc.page_content, "topic": topic} for doc in docs]
        inputs = [{"context": doc_content, "user_rules": instructions}]
        return chain.apply(inputs)
        
    output_file = self.system_directory + 'sop.sop'
    
    analyzed_list = self.load_analyzed_table(self.data_directory, 'analyzed_sop.sop')
    
    #exists_list = self.load_file(self.system_directory, 'categories.sop')
    
    #texts = []
    
    prompt_template_sop = templ

    PROMPT = PromptTemplate(template=prompt_template_sop, input_variables=["context", "user_rules"])

    llm = OpenAI(temperature=0)

    chain = LLMChain(llm=llm, prompt=PROMPT)
        
    count_token = 0
    res_text = ""
    title_exist = False
    #title = ''
    text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(chunk_size=500, chunk_overlap=50)
    
    if (file_ not in analyzed_list):
        #print(analyzed_list)
        print("Загружается файл: ", file_)
        self.debug_log.append("File is loading:" + file_)
        buf = file_.split('.')
        if buf[-1] in ['doc', 'docx'] :
            # Load DOC/DOCX document
            document_txt = ""
            doc = docx.Document(self.data_directory + file_)
            for docpara in doc.paragraphs:
                document_txt += docpara.text + '\n'
            #print(document_txt)
        else:
            # This is a long document we can split up.
            document_txt = self.load_document_text(self.data_directory + file_)
          
        
        with open(output_file, 'w', encoding='utf-8') as file:
            for chunk in text_splitter.split_text(document_txt):
                count_token += self.num_tokens_from_string(str(chunk), "cl100k_base")
                #print(chunk)
                #print('\n ===========================================: ')
                print(count_token)
                buf_res = analyze_chunk(chunk, instructions)[0]["text"]
                
                if not title_exist:
                    lines = res_text.split('\n')
                    for line in lines:
                        print(f'Line: {line}')
                    #buf_res = buf_res.split('\n', 1)[2]
                res_text += buf_res
                #print('\n ===========================================: ')
                #print(f'\n {buf_res["usage"]["total_tokens"]} total tokens used (question-answer). ')
                if not title_exist:
                    title_exist = True
                    #title = res_text.split('\n', 1)[0]
                    print(f'Exist title: True')
                
            file.write(res_text)
        analyzed_list.append(file_)
        #print(document_txt)
        #texts.append(text_splitter.split_text(document_txt))
        '''
        #Сораняем игнор лист (файлы которые уже проанализированы )
        with open(self.data_directory + 'analyzed_sop.sop', 'w', encoding='utf-8') as file:
            for  line in analyzed_list:
                file.write(line + '\n')
        '''
    #res_text = analyze_chunkt(texts[0])[0]["text"]
    #print(res_text)
    #count_token = self.num_tokens_from_string(res_text, "cl100k_base")
    #print(count_token)
    return res_text
  
  def get_input(self, question: str, my_url):
    inputs = {
    "query": question,
    #"requests_result": scrape_site("https://maxximize.me/", 'maxximize.txt')
    "url": my_url,
    }
    return inputs
  
  def get_web_gpt_proposal(self, topic, temp = 0.5, verbose = 0):
  
    # Подготавливаем темплейт для формирования запроса и получения ответа
    #template = self.load_document_text(self.data_directory_prop + self.system_doc_web_prop)
    
    res_f_name = 'magicvaporizers.txt'
    my_url = "https://magicvaporizers.co.uk"
    
    template = """Between >>> and <<< are the raw content text from site.
        Extract the answer to the question '{query}' or say "not found" if the information is not contained.
        Use the format
        Extracted:<answer or "not found">
        >>> {requests_result} <<<
        Extracted:"""
    
    #scrape_site("https://maxximize.me/", 'maxximize.txt')
    #scrape_site(my_url, res_f_name)
    #https://magicvaporizers.co.uk
    
    llm=ChatOpenAI(temperature=0)
    my_chain = load_qa_chain(llm, chain_type="map_reduce")
    #summary_chain = load_summarize_chain(llm, chain_type="map_reduce")
    document_chain = AnalyzeDocumentChain(combine_docs_chain=my_chain)
    
    with open(res_f_name) as f:
        text_doc = f.read()
    
    #print(document_chain.run(input_document=text_doc, question="Remove any irrelevant or duplicate information, correct any errors or inconsistencies, and standardize the data format."))
    #print(document_chain.run(input_document=text_doc, question="Show information about the services provided from the loaded site in the form of a numbered list."))
    #print(document_chain.run(text_doc))
    
    PROMPT = PromptTemplate(
        input_variables=["query", "requests_result"],
        template=template,
    )
    
    # Инициализируем класс
    chain = LLMRequestsChain(llm_chain=LLMChain(llm=llm, prompt=PROMPT))
    
    # Запускаем на вопросе
    #question = self.get_input(topic)
    question = self.get_input("Show information about the services provided from the loaded site in the form of a numbered list.", my_url)
    
    answer = chain(question)
    print(answer)
    return answer  # возвращает ответ
  
  def get_gpt_proposal(self, topic, temp = 0.5, verbose = 0):

    # Выборка документов по схожести с вопросом
    try:
        docs = self.db.similarity_search('Company Maxximaze provides a range of services. Features of our company. Our competencies. Our partners. Our contact details.', k=4)
        message_content = re.sub(r'\n{2}', ' ', '\n '.join([f'\n=====' + doc.page_content + '\n' for i, doc in enumerate(docs)]))
    except:
        message_content = ''
        print('\n No embedding data! \n')
        self.debug_log.append('\n No embedding data! \n')
        
    if verbose:
        print('message_content :\n ======================================== \n' + message_content)

    sys_template = self.chat_manager_system
    
    messages = [
      {"role": "system", "content": sys_template},
      #{"role": "user", "content": f"Analyze the texts of the documents {message_content} and based on the title ( Job title: ) and job description ( Job description: ) write a detailed cover letter for a job offer ( Proposal: ) First of all, you need to follow the system rules. \n{topic}."}
      {"role": "user", "content": f"{topic}\n Additional information : {message_content}"}
      #{"role": "user", "content": topic}
      ]

    #print(f"Prompt: {messages}")
    
    # example token count from the function defined above
    if verbose: 
        self.debug_log.append('\n ===========================================: ')
        self.debug_log.append(f"\n {self.num_tokens_from_messages(messages)} tokens used for the question")
        
    try:
      completion = openai.ChatCompletion.create(
      model=self.model,
      messages=messages,
      temperature=temp
      )
      answer = self.insert_newlines(completion.choices[0].message.content)
      if verbose:
          self.debug_log.append('\n ===========================================: ')
          self.debug_log.append(f'\n {completion["usage"]["total_tokens"]} total tokens used (question-answer). ')
          self.debug_log.append('\n ===========================================: ')
          self.debug_log.append('\n Request price with response :' + str(round(0.002*(completion["usage"]["total_tokens"]/1000), 5)) + ' $ ')
          self.debug_log.append('\n ===========================================: \n')
          #self.debug_log.append('Ответ ChatGPT: ')
          #self.debug_log.append(completion.choices[0].message.content)
    except openai.OpenAIError as e:
      self.debug_log.append(f'OpenAI API Error: {e}')
      answer = f'OpenAI API Error: {e}'
            
    return answer  # возвращает ответ