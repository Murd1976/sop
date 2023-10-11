
from langchain.llms import OpenAI
from langchain.chat_models import ChatOpenAI
from langchain.chains import LLMRequestsChain, LLMChain
from langchain.chains.summarize import load_summarize_chain
from langchain.chains.question_answering import load_qa_chain
from langchain.chains import AnalyzeDocumentChain #, CondenseChain
from langchain.document_loaders import TextLoader
from langchain.docstore.document import Document
import requests
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from langchain.text_splitter import CharacterTextSplitter
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
#import aspose.words as aw

import time
from datetime import datetime
import pathlib
import subprocess
import tempfile

import os, shutil

import re
import logging
logging.getLogger("langchain.text_splitter").setLevel(logging.ERROR) # игнорирование предупреждений
logging.getLogger("chromadb").setLevel(logging.ERROR)
import docx

import openai
import tiktoken
import copy
from dotenv import load_dotenv
from .web_utils import scrape_site

class bcolors:
    HEADER = '\033[95m'
    OKBLUE = '\033[94m'
    OKCYAN = '\033[96m'
    OKGREEN = '\033[92m'
    WARNING = '\033[93m'
    FAIL = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'
	

# Authenticate OpenAI API
load_dotenv()
openai.api_key = os.getenv('OPENAI_API_KEY')



class WorkerОpenAIChat():
  sys_dir = 'ai_sop/db/'
  # папка базы данных Chat
  persist_directory = ''
  # путь к учебным материалам Chat
  data_directory = ''
  db_chat_file = ""
  system_directory = ''
  system_doc_file = 'support_instruction.txt'
  web_directory = ''
  
  __embedding_new =  False

  def __init__(self, em_framework = 'chroma', system_file = '', company_dir = None, mod = 'gpt-3.5-turbo-0301'):
    #global question_history
    self.question_history = []
    self.model = mod
    self.debug_log = []
    self.em_framework = em_framework
    
    self.persist_directory = self.sys_dir + 'chat/embedding/' + company_dir
    self.system_directory = self.sys_dir + 'chat/sys/' + company_dir
    self.system_doc_file = system_file
    self.data_directory = self.sys_dir + 'data/' + company_dir
    self.web_directory = self.sys_dir + 'chat/web/' + company_dir
    
    if not os.path.exists(self.persist_directory):
        os.mkdir(self.persist_directory)
    if not os.path.exists(self.system_directory):
        os.mkdir(self.system_directory)
    if not os.path.exists(self.data_directory):
        os.mkdir(self.data_directory)
    if not os.path.exists(self.web_directory):
        os.mkdir(self.web_directory)
    
    self.chat_manager_system = self.get_text_from(self.system_directory, self.system_doc_file)
    
    if em_framework == 'chroma':
        self.persist_directory = self.persist_directory + 'chroma/'
        if not os.path.exists(self.persist_directory):
            os.mkdir(self.persist_directory)
        # Если База данных embedding уже создана ранее
        if os.path.exists(self.persist_directory + 'embedding_info.inf'):
            print("We use a ready-made database. Path: ", self.persist_directory)
            self.debug_log.append("\n We use a ready-made database. Path: " + self.persist_directory)
            self.db = Chroma(persist_directory=self.persist_directory,
                                embedding_function=OpenAIEmbeddings())
        else:
            print("Embeddings database not found!")
            self.debug_log.append("\n Embeddings database not found!")
            self.db = Chroma(persist_directory=self.persist_directory, embedding_function=OpenAIEmbeddings())

    if em_framework == 'faiss':
        self.persist_directory = self.persist_directory + 'faiss/'
        # Load index from disk
        self.db = FAISS.load_local(self.persist_directory + 'faiss_index') 

  # def get_key(self):
  #   openai.api_key = getpass.getpass(prompt='Введите секретный ключ для сервиса chatGPT: ')
  #   import os
  #   os.environ["OPENAI_API_KEY"] = openai.api_key

  def set_key(self):
      password_input = widgets.Password(
          description='Enter password:',
          layout=widgets.Layout(width='500px'),
          style={'description_width': 'initial', 'white-space': 'pre-wrap', 'overflow': 'auto'})
      login_button = widgets.Button(description='Авторизация')
      output = widgets.Output()

      def on_button_clicked(_):
          with output:
              openai.api_key = password_input.value
              os.environ["OPENAI_API_KEY"] = openai.api_key
              print(f'{bcolors.OKGREEN}{bcolors.BOLD}Ключ сохранен!{bcolors.ENDC}')
              password_input.layout.display = 'none'
              login_button.layout.display = 'none'

      login_button.on_click(on_button_clicked)
      display(widgets.VBox([password_input, login_button, output]))

  @property  
  def em_mode(self):
    return self.__embedding_new
  
  @em_mode.setter  
  def em_mode(self, mode):
    if mode == '1':
        self.__embedding_new = False
    else:
        self.__embedding_new = True

  def load_document_text(self, url: str) -> str:
      with open(url, "r") as f:
        text = f.read()

      return text
      
  def file_list(self, f_path):
                
    f_list = os.listdir(f_path)
    f_list = [ff for ff in f_list if (os.path.isfile(f_path + ff) and ('~' not in ff))]
        
    return f_list

  def del_all_in_dir(self, dir_path):
    if os.path.exists(dir_path):
        for files in os.listdir(dir_path):
            path = os.path.join(dir_path, files)
            try:
                shutil.rmtree(path)
            except OSError:
                os.remove(path)
                
  def get_text_from(self, doc_dir, file_):
    document_txt = ""
    if file_:
        with open(doc_dir + file_, "r") as f:                
            self.debug_log.append("File is loading:" + file_)
            buf = file_.split('.')
            if buf[-1] in ['doc', 'docx'] :
                # Load DOC/DOCX document
                
                try:
                    doc = docx.Document(doc_dir + file_)
                    for docpara in doc.paragraphs:
                        document_txt += docpara.text + '\n'
                except Exception as e:
                    self.debug_log.append(f"Error processing file {file_}: {str(e)}")
                    
            elif buf[-1] == 'txt':
                try:
                    document_txt = self.load_document_text(doc_dir + file_)                
                    self.debug_log.append("File is loading:" + doc_dir + file_)
                except Exception as e:
                    self.debug_log.append(f"Error processing file {file_}: {str(e)}")
            
            
    return document_txt

  def num_tokens_from_string(self, string: str, encoding_name: str) -> int:
      """Returns the number of tokens in a text string."""
      encoding = tiktoken.get_encoding(encoding_name)
      num_tokens = len(encoding.encode(string))
      return num_tokens
      
      
    
  def create_embedding(self, doc_dir="", un_list = []):
    if self.em_framework == 'chroma':
        used_list = self.create_embedding_chroma(doc_dir= doc_dir, un_list = un_list)
    elif self.em_framework == 'faiss':    
        used_list = self.create_embedding_faiss(doc_dir= doc_dir, unlist = un_list)
        
    return used_list
        
  def create_embedding_chroma(self, doc_dir="", un_list = []):
        
    self.source_chunks = []
    buf_chunks = []
    splitter = RecursiveCharacterTextSplitter(["<Chunk>", '\n\n', '\n', ' '], chunk_size=1024, chunk_overlap=0)
    print(f'Unused list: {un_list}')
    used_list = []
    if self.__embedding_new :
        print('Creating new embeddings!')
        
        self.del_all_in_dir(self.persist_directory)
        self.debug_log.append('\n Previouse embeddings deleted!')
        self.db = Chroma(persist_directory=self.persist_directory, embedding_function=OpenAIEmbeddings())
    else:        
        print('Adding embeddings!')
   
    
    # проходимся по всем данным
    count_em_token = 0 # all tokens gones thrue embeddings
    f_info = open(self.persist_directory + 'embedding_info.inf', 'w')
    f_info.write('Used files: \n')
    for file_ in sorted(self.file_list(doc_dir)): #go thrue all files in doc_dir
        #if not self.__embedding_new:
        if file_ not in un_list:
            continue
            
        print("Загружается файл: ", file_)
        
        self.debug_log.append("\n File is loading:" + doc_dir)
        # разбиваем на несколько частей с помощью метода split_text
        if (os.path.isfile(doc_dir + file_)):
            
            try:
                document_txt = self.get_text_from(doc_dir, file_) #get text from current file
                f_info.write(str(file_) + '\n')
            except Exception as e:
                self.debug_log.append(f"\n Error processing file {file_}: {str(e)}") 
                
            count_token = 0 #counter for OpenAI (tokens limit 140K per minute )    
            for chunk in splitter.split_text(document_txt):
                #print('Длина символов =  ', len(chunk))

                count_token += self.num_tokens_from_string(chunk, "cl100k_base")
                
                if count_token > 140000:
                   
                    #print('bCount: ', count_token, ' Tokens:  ', self.num_tokens_from_string(' '.join([x.page_content for x in buf_chunks]), "cl100k_base"))
                  
                    count_token = 0
                    self.source_chunks.append(copy.deepcopy(buf_chunks))
                   
                    #print('bSize: ', len(buf_chunks), '\n')
                    buf_chunks.clear()
                   
                    buf_chunks.append(Document(page_content=chunk, metadata={'source': file_}))
                else:
                    buf_chunks.append(Document(page_content=chunk, metadata={'source': file_}))
                    
            #print('Count: ', count_token, ' Tokens:  ', self.num_tokens_from_string(' '.join([x.page_content for x in buf_chunks]), "cl100k_base"))        
            #print('Size: ', len(buf_chunks), '\n')            
            self.source_chunks.append(copy.deepcopy(buf_chunks))
            buf_chunks.clear()
                
            count_token = 0       
            for i in range(len(self.source_chunks)):
                n_buf = self.num_tokens_from_string(' '.join([x.page_content for x in self.source_chunks[i]]), "cl100k_base")
                count_token += n_buf
                if count_token > 140000:
                    count_token = n_buf
                    #print('\n PAUSE \n')
                    time.sleep(77)
                self.db.add_documents(documents=self.source_chunks[i])
                #print('sSize: ', len(self.source_chunks[i]))
                self.debug_log.append('\n sSize: ' + str(len(self.source_chunks[i])))
                count_em_token += n_buf                
                #print(f'Curr 140K chunk: {i} Counter: Global- {count_em_token} Curr file- {n_buf} ')
                self.debug_log.append(f'Block of current file: {i} Token counter: Global- {count_em_token} Current block- {n_buf} ')
                
                  
            self.db.persist()
            used_list.append(file_)
            self.source_chunks.clear()
                
    #print('Count: ', count_token, ' Tokens:  ', num_tokens_from_string(' '.join([x.page_content for x in buf_chunks]), "cl100k_base"))
    #print('Size: ', len(buf_chunks), '\n')
    
    #print('G_Size: ', self.source_chunks[0])
    #print('G_Size: ', self.source_chunks[1])
    #print('G_Size: ', self.source_chunks[2])
       
    f_info.close()
    
    print('\n ===========================================:')
    print('\n Number of tokens in source document: ', count_em_token)
    print('\n Request price: ', 0.0004*(count_em_token/1000), ' $')
    print('\n ===========================================')
    
    self.debug_log.append('\n ===========================================: ')
    self.debug_log.append('\n Number of tokens in source document: ' + str(count_em_token))
    self.debug_log.append('\n Request price: ' + str(round(0.0004*(count_em_token/1000), 5)) + ' $')
    self.debug_log.append('\n =========================================== \n')
   
    return used_list

  def create_embedding_faiss(self, doc_dir="", u_list = []):
        
    self.source_chunks = []
    self.buf_chunks = []

  def answer(self, system, topic, temp = 1):
      messages = [
        {"role": "system", "content": system},
        {"role": "user", "content": topic}
        ]

      completion = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages,
        temperature=temp
        )

      return completion.choices[0].message.content

  def num_tokens_from_messages(self, messages):
      """
      Возвращает количество токенов, используемых списком сообщений.
      """
      try:
          # Пытаемся получить кодировку для выбранной модели
          encoding = tiktoken.encoding_for_model(self.model)
      except KeyError:
          # Если кодировка для выбранной модели не найдена, используем кодировку "cl100k_base"
          encoding = tiktoken.get_encoding("cl100k_base")
      # Если выбранная модель это "gpt-3.5-turbo-0301"
      if self.model == "gpt-3.5-turbo-0301":
          # Инициализируем счетчик токенов
          num_tokens = 0
          # Проходимся по каждому сообщению в списке сообщений
          for message in messages:
              # Каждое сообщение обрамляется токенами <im_start> и <im_end>, а также символами новой строки, всего 4 токена
              num_tokens += 4
              # Проходимся по каждому полю в сообщении (ключ и значение)
              for key, value in message.items():
                  # Считаем количество токенов в значении и добавляем их в счетчик токенов
                  num_tokens += len(encoding.encode(value))
                  # Если ключ это "name", то это означает что роль (role) опущена
                  if key == "name":
                      # Роль всегда требуется и всегда занимает 1 токен, так что вычитаем 1 из счетчика
                      num_tokens += -1
          # Каждый ответ начинается с токена <im_start>assistant, так что добавляем 2 в счетчик
          num_tokens += 2
          # Возвращаем количество токенов
          return num_tokens
      else:
          # Если выбранная модель не поддерживается, генерируем исключение
          raise NotImplementedError(f"""num_tokens_from_messages() is not presently implemented for model {self.model}.
  See https://github.com/openai/openai-python/blob/main/chatml.md for information on how messages are converted to tokens.""")


  def insert_newlines(self, text: str, max_len: int = 170) -> str:
      words = text.split()
      lines = []
      current_line = ""
      for word in words:
          if len(current_line + " " + word) > max_len:
              lines.append(current_line)
              current_line = ""
          current_line += " " + word
      lines.append(current_line)
      return "\n".join(lines)


  def dialog(self):
      user = ''
      dialog = ''


      print(f'{bcolors.OKBLUE}{bcolors.BOLD}С чем связан ваш интерес к искусственному интеллекту?{bcolors.ENDC}')

      while user.lower() not in ['stop', 'exit', 'выход']:
          user = input('User: ')
          if user == 'Stop': break

          dialog += '\n\n' + 'User: ' + user
          add_dialog = self.answer(expert_prompt, user)

          dialog += '\n\n' + 'Assistant: ' + add_dialog
          print(f'\n{bcolors.OKBLUE}{bcolors.BOLD}Assistant:{bcolors.ENDC} {self.insert_newlines(add_dialog)}')
          report = self.answer(validation_prompt, dialog)
          answer_text = self.answer(action_prompt, report)

          print(f'\n{bcolors.OKGREEN}{bcolors.BOLD}System report:\n {bcolors.ENDC}{report}')
          print(f'\n{bcolors.HEADER}{bcolors.BOLD}Assistant: {bcolors.ENDC}{self.insert_newlines(answer_text)}\n\n')

      return dialog

  def answer_index(self, topic, user_question, temp=1, verbose=0):

      # Selecting documents similar to the question
      docs = self.search_index.similarity_search(f'Subject": {user_question["subject"]}\n"' + "\n\n Сurrent issue: " + user_question['issue'], k=5)
      #docs = self.search_index.similarity_search(topic, k=5)
      
      #docs = self.search_index.get_relevant_documents(topic)
      
      #print(f'\n {topic} \n')
      message_content = re.sub(r'\n{2}', ' ', '\n '.join([f'\nText №{i+1}\n=====================' + doc.page_content + '\n' for i, doc in enumerate(docs)]))
      if verbose: 
        #print('\n ===========================================: ')
        print('message_content :\n ======================================== \n', message_content)
        
        #self.debug_log.append('\n ===========================================: \n')
        #self.debug_log.append('message_content :\n ======================================== \n' + message_content)
     
      print(f'\n System: {self.chat_manager_system} \n')
      messages = [
          {"role": "system", "content": self.chat_manager_system + f"{message_content}"},
          {"role": "user", "content": topic}
      ]

      if verbose: 
        #print('\n ===========================================: ')
        #print(f"{self.num_tokens_from_messages(messages, 'gpt-3.5-turbo-0301')} tokens used for the question")
        
        self.debug_log.append('\n ===========================================: ')
        self.debug_log.append(f"\n {self.num_tokens_from_messages(messages)} tokens used for the question")

      completion = openai.ChatCompletion.create(
          model="gpt-3.5-turbo",
          messages=messages,
          temperature=temp
      )

      if verbose:
        '''
        print('\n ===========================================: ')
        print(f'{completion["usage"]["total_tokens"]} total tokens used (question-answer).')
        print('\n ===========================================: ')
        print('Request price with response :', 0.002*(completion["usage"]["total_tokens"]/1000), ' $')
        print('============================================: \n')
        '''
        self.debug_log.append('\n ===========================================: ')
        self.debug_log.append(f'\n {completion["usage"]["total_tokens"]} total tokens used (question-answer). ')
        self.debug_log.append('\n ===========================================: ')
        self.debug_log.append('\n Request price with response :' + str(round(0.002*(completion["usage"]["total_tokens"]/1000), 5)) + ' $ ')
        self.debug_log.append('\n ===========================================: \n')
        
      answer = completion.choices[0].message.content #self.insert_newlines(completion.choices[0].message.content)
      #print('ANSWER : \n', answer)
      return answer  # возвращает ответ

  def num_tokens_from_history(self, messages, model="gpt-3.5-turbo-0301"):
      """Returns the number of tokens used by a list of messages."""
      try:
          encoding = tiktoken.encoding_for_model(model)
      except KeyError:
          encoding = tiktoken.get_encoding("cl100k_base")
      if model == "gpt-3.5-turbo-0301":  # note: future models may deviate from this
          num_tokens = 0
          for message in messages:
              question, answer = message  # распаковываем кортеж
              num_tokens += 4  # every message follows <im_start>{role/name}\n{content}<im_end>\n
              num_tokens += len(encoding.encode(question))
              num_tokens += len(encoding.encode(answer))
          num_tokens += 2  # every reply is primed with <im_start>assistant
          return num_tokens
      else:
          raise NotImplementedError(f"""num_tokens_from_messages() is not presently implemented for model {model}.""")

  def summarize_questions(self, dialog):
      # Применяем модель GPT-3 для суммаризации вопросов
      messages = [
          {"role": "system", "content": "You are an AI-powered technical support assistant. You are able to professionally summarize the dialogues sent to you by the manager and the client. Your task is to summarize the dialogue that came to you."},
          {"role": "user", "content": "Summarize the following dialogue between a sales manager and a customer: " + " ".join(dialog)}
      ]

      completion = openai.ChatCompletion.create(
          model="gpt-3.5-turbo",
          messages=messages,
          temperature=0.3,  # Используем более низкую температуру для более определенной суммаризации
          max_tokens=500  # Ограничиваем количество токенов для суммаризации
      )

      return completion.choices[0].message.content

  def answer_user_question(self, user_question, temp = 0.5, verbose = 0) -> str:

      #knowledge_base_text = load_document_text(knowledge_base_url)

      #print('\n', knowledge_base_text)
      
      # Создаем индексы поиска
      #knowledge_base_index = create_search_index(knowledge_base_text)

      # Если в истории более одного вопроса, применяем суммаризацию
      summarized_history = ""
      if len(self.question_history) > 0:
          summarized_history = "Here is a summary of the previous dialogue: " + self.summarize_questions([q + ' ' + (a if a is not None else '') for q, a in self.question_history])

      # Добавляем явное разделение между историей диалога и текущим вопросом
      #query = '"prompt":" Company: PayStabs\n User name: Jenny021914\n Subject: PSC - Contact Form\n' + query + '"'
      input_text =summarized_history + f'"prompt":" Company: {user_question["company"]}\n User name: {user_question["user_name"]}\n Subject": {user_question["subject"]}\n"' + "\n\n Сurrent issue: " + user_question['issue']

      # Извлечение наиболее похожих отрезков текста из базы знаний и получение ответа модели
      answer_text = self.answer_index(input_text, user_question, temp=temp, verbose = verbose)
      '''
      # Добавляем вопрос пользователя и ответ системы в историю
      tt = datetime.now().strftime('%m/%d/%Y %H:%M:%S')
      self.question_history.append((tt + '\n' + user_question['issue'], answer_text if answer_text is not None else ''))

      # Выводим суммаризированный текст, который видит модель
      if summarized_history != "":
          #print("Here is the summarized text that the model sees:\n", summarized_history)
          self.debug_log.append("Here is the summarized text that the model sees:\n" + summarized_history)
      #print('\n', user_question, '\n')
      self.debug_log.append('\n' + user_question['issue'] + '\n')
      '''
      return answer_text

  def run_dialog(self, system_doc_url, knowledge_base_url):
      question_history = []
      dialog = ""
      while True:
          user_question = input('User: ')
          if user_question.lower() == 'stop':
              break
          answer = self.answer_user_question_dialog(system_doc_url, knowledge_base_url, user_question, question_history)
          dialog += f'\n User: {user_question} \n Assistant: {answer}'
          print('\n Assistant: ', answer)

      return

#========================================================================================================================================================================
#========================================================================================================================================================================

class WorkerОpenAIProposal():
  sys_dir = 'ai_sop/db/'
  # папка базы данных Proposal
  persist_directory = ''
  # путь к учебным материалам Proposal
  data_directory = ''
  #db_proposal_file = 'train_data_jobs.jsonl'
  system_directory = ''
  system_doc_file = ''
  web_directory = ''
  
  system_doc_web_prop = 'html_proposal_instruction.txt'
  

  def __init__(self, system_file = '', company_dir = None, mod = 'gpt-3.5-turbo-0301'):
    self.model = mod
    self.debug_log = []
    self.persist_directory = self.sys_dir + 'proposal/embedding/' + company_dir
    self.system_directory = self.sys_dir + 'proposal/sys/' + company_dir
    self.system_doc_file = system_file
    self.data_directory = self.sys_dir + 'data/' + company_dir
    self.web_directory = self.sys_dir + 'proposal/web/' + company_dir
    
    if not os.path.exists(self.persist_directory):
        os.mkdir(self.persist_directory)
    if not os.path.exists(self.system_directory):
        os.mkdir(self.system_directory)
    if not os.path.exists(self.data_directory):
        os.mkdir(self.data_directory)
    if not os.path.exists(self.web_directory):
        os.mkdir(self.web_directory)
    
    '''
    if chat_manager_system:
      self.chat_manager_system = self.load_document_text(chat_manager_system)
    else:
      self.chat_manager_system = " "
    '''
    try:
      self.chat_manager_system = self.load_document_text(self.system_directory + self.system_doc_file)
    except:
      self.chat_manager_system = " "
      
    # Если База данных embedding уже создана ранее
    if os.path.exists(self.persist_directory + 'embedding_info.inf'):
      print("We use a ready-made database. Path: ", self.persist_directory)
      self.debug_log.append("We use a ready-made database. Path: " + self.persist_directory)
      self.db = Chroma(persist_directory=self.persist_directory,
                            embedding_function=OpenAIEmbeddings())
    else:
      print("Embeddings database not found!")
      self.debug_log.append("Embeddings database not found!")
      #self.db = Chroma(persist_directory=self.persist_directory, embedding_function=OpenAIEmbeddings())

  # def get_key(self):
  #   openai.api_key = getpass.getpass(prompt='Введите секретный ключ для сервиса chatGPT: ')
  #   import os
  #   os.environ["OPENAI_API_KEY"] = openai.api_key

  def load_document_text(self, url: str) -> str:
      with open(url, "r") as f:
        text = f.read()

      return text

  def num_tokens_from_string(self, string: str, encoding_name: str) -> int:
      """Returns the number of tokens in a text string."""
      encoding = tiktoken.get_encoding(encoding_name)
      num_tokens = len(encoding.encode(string))
      return num_tokens

  def create_embedding(self, doc_dir="", persist_directory=""):
    

    self.source_chunks = []
    self.buf_chunks = []
    splitter = RecursiveCharacterTextSplitter(["<Chunk>", '\n\n', '\n', ' '], chunk_size=1024, chunk_overlap=0)

    
    #print('Files: ', os.listdir(doc_dir))

    #print("File is loading: ", doc_dir)
    
    # проходимся по всем данным
    count_token = 0
    for file_ in sorted(os.listdir(doc_dir)):
        print("Загружается файл: ", file_)
        self.debug_log.append("File is loading:" + doc_dir)
        # разбиваем на несколько частей с помощью метода split_text
        
        with open(doc_dir + file_, "r") as f:
          for chunk in splitter.split_text(f.read()):
              #print('Длина символов =  ', len(chunk))

              count_token += self.num_tokens_from_string(chunk, "cl100k_base")
              if count_token > 140000:
               
                #print('Count: ', count_token, ' Tokens:  ', num_tokens_from_string(' '.join([x.page_content for x in self.buf_chunks]), "cl100k_base"))
                
                count_token = 0
                self.source_chunks.append(copy.deepcopy(self.buf_chunks))
               
                #print('Size: ', len(self.buf_chunks), '\n')
                self.buf_chunks.clear()
                
                self.buf_chunks.append(Document(page_content=chunk, metadata={'source': file_}))
              else:
                self.buf_chunks.append(Document(page_content=chunk, metadata={'source': file_}))

    self.source_chunks.append(copy.deepcopy(self.buf_chunks))

    self.db = Chroma(persist_directory=persist_directory, embedding_function=OpenAIEmbeddings())
    #print('Count: ', count_token, ' Tokens:  ', num_tokens_from_string(' '.join([x.page_content for x in buf_chunks]), "cl100k_base"))
    #print('Size: ', len(buf_chunks), '\n')
    count_token = 0
    #print('G_Size: ', self.source_chunks[0])
    #print('G_Size: ', self.source_chunks[1])
    #print('G_Size: ', self.source_chunks[2])
    for i in range(len(self.source_chunks)):
      self.db.add_documents(documents=self.source_chunks[i])
      print('sSize: ', len(self.source_chunks[i]))
      self.debug_log.append('sSize: ' + str(len(self.source_chunks[i])))
      count_token += self.num_tokens_from_string(' '.join([x.page_content for x in self.source_chunks[i]]), "cl100k_base")
      print(i, 'Counter: ', count_token)
      self.debug_log.append(str(i) + 'Counter: ' + str(count_token))
      time.sleep(77)
      
    self.db.persist()
    
    f = open(persist_directory + 'embedding_info.inf', 'w')
    f.write(str(datetime.now()))
    f.close()
    
    print('\n ===========================================:')
    print('\n Number of tokens in source document: ', count_token)
    print('\n Request price: ', 0.0004*(count_token/1000), ' $')
    print('\n ===========================================')
    
    self.debug_log.append('\n ===========================================: ')
    self.debug_log.append('\n Number of tokens in source document: ' + str(count_token))
    self.debug_log.append('\n Request price: ' + str(0.0004*(count_token/1000)) + ' $')
    self.debug_log.append('\n =========================================== \n')


  def insert_newlines(self, text: str, max_len: int = 170) -> str:
      words = text.split()
      lines = []
      current_line = ""
      for word in words:
          if len(current_line + " " + word) > max_len:
              lines.append(current_line)
              current_line = ""
          current_line += " " + word
      lines.append(current_line)
      return "\n".join(lines)
      
  # пример подсчета токенов
  def num_tokens_from_messages(self, messages):
      """Returns the number of tokens used by a list of messages."""
      try:
          encoding = tiktoken.encoding_for_model(self.model)
      except KeyError:
          encoding = tiktoken.get_encoding("cl100k_base")
      if self.model == "gpt-3.5-turbo-0301":  # note: future models may deviate from this
          num_tokens = 0
          for message in messages:
              num_tokens += 4  # every message follows <im_start>{role/name}\n{content}<im_end>\n
              for key, value in message.items():
                  num_tokens += len(encoding.encode(value))
                  if key == "name":  # if there's a name, the role is omitted
                      num_tokens += -1  # role is always required and always 1 token
          num_tokens += 2  # every reply is primed with <im_start>assistant
          return num_tokens
      else:
          raise NotImplementedError(f"""num_tokens_from_messages() is not presently implemented for model {model}.
  See https://github.com/openai/openai-python/blob/main/chatml.md for information on how messages are converted to tokens.""")
  
  
  def load_file(self, f_dir, f_name):
    
    text = ""
    try:
        with open(f_dir + f_name, "r") as f:
            text = f.read()
        f.close()
        print("File is loading:" + f_dir + f_name)
        self.debug_log.append("File is loading:" + f_dir + f_name)
    except:
        
        print("File not found:" + f_dir + f_name)
        self.debug_log.append("File not found:" + f_dir + f_name)
        
    return text
    
  def load_analyzed_table(self, f_dir, f_name):
    res_list = []
    try:
        with open(f_dir + f_name, "r") as f:
            text = f.read()
        f.close()
        for line in data.splitlines():
            res_list.append(line)
        print("File is loading:" + f_dir + f_name)
        self.debug_log.append("File is loading:" + f_dir + f_name)
    except:
        res_list.append(f_name)
        print("File not found:" + f_dir + f_name)
        self.debug_log.append("File not found:" + f_dir + f_name)
        
    return res_list
    
  
   
    
  def get_gpt_proposal(self, topic, temp = 0.5, verbose = 0):

    # Выборка документов по схожести с вопросом
    try:
        docs = self.db.similarity_search('Company Maxximaze provides a range of services. Features of our company. Our competencies. Our partners. Our contact details.', k=4)
        message_content = re.sub(r'\n{2}', ' ', '\n '.join([f'\n=====' + doc.page_content + '\n' for i, doc in enumerate(docs)]))
    except:
        message_content = ''
        print('\n No embedding data! \n')
        self.debug_log.append('\n No embedding data! \n')
        
    if verbose:
        print('message_content :\n ======================================== \n' + message_content)

    sys_template = self.chat_manager_system
    
    messages = [
      {"role": "system", "content": sys_template},
      #{"role": "user", "content": f"Analyze the texts of the documents {message_content} and based on the title ( Job title: ) and job description ( Job description: ) write a detailed cover letter for a job offer ( Proposal: ) First of all, you need to follow the system rules. \n{topic}."}
      {"role": "user", "content": f"{topic}\n Additional information : {message_content}"}
      #{"role": "user", "content": topic}
      ]

    print(f"Prompt: {messages}")
    
    # example token count from the function defined above
    if verbose: 
        self.debug_log.append('\n ===========================================: ')
        self.debug_log.append(f"\n {self.num_tokens_from_messages(messages)} tokens used for the question")
        
    try:
      completion = openai.ChatCompletion.create(
      model=self.model,
      messages=messages,
      temperature=temp
      )
      answer = self.insert_newlines(completion.choices[0].message.content)
      if verbose:
          self.debug_log.append('\n ===========================================: ')
          self.debug_log.append(f'\n {completion["usage"]["total_tokens"]} total tokens used (question-answer). ')
          self.debug_log.append('\n ===========================================: ')
          self.debug_log.append('\n Request price with response :' + str(round(0.002*(completion["usage"]["total_tokens"]/1000), 5)) + ' $ ')
          self.debug_log.append('\n ===========================================: \n')
          #self.debug_log.append('Ответ ChatGPT: ')
          #self.debug_log.append(completion.choices[0].message.content)
    except openai.OpenAIError as e:
      self.debug_log.append(f'OpenAI API Error: {e}')
      answer = f'OpenAI API Error: {e}'
            
    return answer  # возвращает ответ